#!/usr/bin/env python3
"""将 MD 变更文件转换为 docx_revision JSON 配置"""
import argparse
import json
import re
import sys
from pathlib import Path


def check_ambiguity(edit, doc):
    """检查文本在段落中是否出现多次"""
    if edit["type"] not in ("replace", "delete"):
        return None

    para_idx = edit["paragraph_index"]
    if para_idx >= len(doc.paragraphs):
        return f"段落索引 {para_idx} 超出范围"

    para_text = doc.paragraphs[para_idx].text
    old_text = edit.get("old_text") or edit.get("text")

    if old_text not in para_text:
        return f"文本 '{old_text[:30]}...' 未在段落 {para_idx} 中找到"

    count = para_text.count(old_text)
    if count > 1:
        positions = []
        start = 0
        while True:
            pos = para_text.find(old_text, start)
            if pos == -1:
                break
            positions.append((pos, pos + len(old_text)))
            start = pos + 1

        pos_str = "\n".join([f"  位置{i+1}: 第{p[0]}-{p[1]}字符" for i, p in enumerate(positions)])
        return f"⚠️ 歧义检测: \"{old_text[:30]}...\" 在段落 {para_idx} 出现{count}次\n{pos_str}\n\n请在 MD 中添加位置信息:\n  删除: \"{old_text}\" (第{positions[0][0]}-{positions[0][1]}字符)"

    return None

def parse_frontmatter(content):
    """解析 YAML frontmatter"""
    match = re.match(r'^---\n(.*?)\n---\n', content, re.DOTALL)
    if not match:
        raise ValueError("缺少 YAML frontmatter")
    
    frontmatter = {}
    for line in match.group(1).split('\n'):
        if ':' in line:
            key, value = line.split(':', 1)
            frontmatter[key.strip()] = value.strip()
    
    return frontmatter, content[match.end():]

def parse_sections(body):
    """解析各 section"""
    sections = {}
    current_section = None
    
    for line in body.split('\n'):
        if line.startswith('# '):
            current_section = line[2:].strip()
            sections[current_section] = []
        elif current_section and line.strip() and line.strip() != '---':
            sections[current_section].append(line)
    
    return sections

def parse_comments(lines):
    """解析 Comments section"""
    comments = []
    current = None
    
    for line in lines:
        # 匹配 "Para N: title" 或 "Para N: title (范围)"
        match = re.match(r'^Para\s+(\d+):\s*(.+)', line)
        if match:
            if current:
                comments.append(current)
            current = {
                "paragraph_index": int(match.group(1)),
                "text": "",
                "author": None,
                "initials": None,
                "start_pos": None,
                "end_pos": None
            }
            continue
        
        # 匹配范围信息
        range_match = re.match(r'^>\s*选中范围:\s*第(\d+)-(\d+)字符', line)
        if range_match and current:
            current["start_pos"] = int(range_match.group(1))
            current["end_pos"] = int(range_match.group(2))
            continue
        
        # 匹配缩写
        initials_match = re.match(r'^>\s*缩写:\s*(\w+)', line)
        if initials_match and current:
            current["initials"] = initials_match.group(1)
            continue
        
        # 普通文本行
        if current and line.strip() and not line.startswith('>'):
            if current["text"]:
                current["text"] += "\n" + line.strip()
            else:
                current["text"] = line.strip()
    
    if current:
        comments.append(current)
    
    return comments


def parse_text_edits(lines):
    """解析 Text Edits section"""
    edits = []
    current_para = None
    current_title = None
    
    for line in lines:
        # 匹配 "Para N: title"
        para_match = re.match(r'^Para\s+(\d+):\s*(.+)', line)
        if para_match:
            current_para = int(para_match.group(1))
            current_title = para_match.group(2)
            continue
        
        if current_para is None:
            continue
        
        # 替换: 将 "old" 改为 "new"
        replace_match = re.match(r'^将\s*"(.+?)"\s*改为\s*"(.+?)"', line)
        if replace_match:
            edits.append({
                "type": "replace",
                "paragraph_index": current_para,
                "old_text": replace_match.group(1),
                "new_text": replace_match.group(2),
                "author": None
            })
            continue
        
        # 插入: 在开头/末尾插入: text
        insert_match = re.match(r'^在(开头|末尾)插入:\s*(.+)', line)
        if insert_match:
            position = 0 if insert_match.group(1) == '开头' else None
            edits.append({
                "type": "insert",
                "paragraph_index": current_para,
                "text": insert_match.group(2),
                "position": position,
                "author": None
            })
            continue
        
        # 删除: 删除: "text" 或 删除: "text" (第N-M字符)
        delete_match = re.match(r'^删除:\s*"(.+?)"(?:\s*\(第(\d+)-(\d+)字符\))?', line)
        if delete_match:
            edit = {
                "type": "delete",
                "paragraph_index": current_para,
                "text": delete_match.group(1),
                "author": None
            }
            if delete_match.group(2):
                edit["start_pos"] = int(delete_match.group(2))
                edit["end_pos"] = int(delete_match.group(3))
            edits.append(edit)
            continue
    
    return edits


FORMAT_KEYWORDS = {
    '居中对齐': ('alignment', 'center'),
    '左对齐': ('alignment', 'left'),
    '右对齐': ('alignment', 'right'),
    '两端对齐': ('alignment', 'justify'),
    '加粗': ('bold', True),
    '斜体': ('italic', True),
}


def parse_format_value(text):
    """解析格式值"""
    match = re.match(r'行距([\d.]+)倍', text)
    if match:
        return ('line_spacing', float(match.group(1)))

    match = re.match(r'段前([\d.]+)pt', text)
    if match:
        return ('space_before', float(match.group(1)))

    match = re.match(r'段后([\d.]+)pt', text)
    if match:
        return ('space_after', float(match.group(1)))

    match = re.match(r'左缩进([\d.]+)pt', text)
    if match:
        return ('indent_left', float(match.group(1)))

    match = re.match(r'右缩进([\d.]+)pt', text)
    if match:
        return ('indent_right', float(match.group(1)))

    match = re.match(r'字号([\d.]+)pt', text)
    if match:
        return ('font_size', float(match.group(1)))

    return None


def parse_format_edits(lines):
    """解析 Format Edits section"""
    edits = []
    current_para = None
    current_changes = {}

    for line in lines:
        para_match = re.match(r'^Para\s+(\d+)(?:-(\d+))?:\s*(.*)', line)
        if para_match:
            if current_para is not None and current_changes:
                edits.append({
                    "scope": "paragraph",
                    "paragraph_index": current_para,
                    "changes": current_changes.copy(),
                    "author": None
                })
            current_para = int(para_match.group(1))
            current_changes = {}
            continue

        if current_para is None:
            continue

        items = [x.strip() for x in line.split(',')]
        for item in items:
            for keyword, (key, value) in FORMAT_KEYWORDS.items():
                if keyword in item:
                    current_changes[key] = value
                    break
            else:
                result = parse_format_value(item)
                if result:
                    current_changes[result[0]] = result[1]

    if current_para is not None and current_changes:
        edits.append({
            "scope": "paragraph",
            "paragraph_index": current_para,
            "changes": current_changes,
            "author": None
        })

    return edits


def parse_table_edits(lines):
    """解析 Table Edits section"""
    edits = []
    current_table = None
    
    cn_num = {'一':1, '二':2, '三':3, '四':4, '五':5, '六':6, '七':7, '八':8, '九':9, '十':10}
    
    for line in lines:
        table_match = re.match(r'^表格(\d+):', line)
        if table_match:
            current_table = int(table_match.group(1)) - 1  # 转为0-based
            continue
        
        if current_table is None:
            continue
        
        # 第N行下方加一行
        insert_match = re.match(r'^\s*第(\d+)行下方加一行', line)
        if insert_match:
            edits.append({
                "type": "insert_row",
                "table_index": current_table,
                "row_index": int(insert_match.group(1)),
                "author": None
            })
            continue
        
        # 删掉第N行
        delete_match = re.match(r'^\s*删掉第(\d+)行', line)
        if delete_match:
            edits.append({
                "type": "delete_row",
                "table_index": current_table,
                "row_index": int(delete_match.group(1)),
                "author": None
            })
            continue
        
        # 合并第N行的X-Y列 (字母列)
        merge_match = re.match(r'^\s*合并第([一二三四五六七八九十\d]+)行的第?([A-Z])-([A-Z])列', line)
        if merge_match:
            row_str = merge_match.group(1)
            row = (cn_num[row_str] if row_str in cn_num else int(row_str)) - 1
            start_col = ord(merge_match.group(2)) - ord('A')
            end_col = ord(merge_match.group(3)) - ord('A')
            edits.append({
                "type": "merge_cells",
                "table_index": current_table,
                "start_cell": [row, start_col],
                "end_cell": [row, end_col],
                "author": None
            })
            continue
        
        # 合并第N行的X个格子 (自然语言/中文数字)
        merge_natural = re.match(r'^\s*合并第?([一二三四五六七八九十\d]+)行的([一二三四五六七八九十\d]+)个格子', line)
        if merge_natural:
            row_str = merge_natural.group(1)
            count_str = merge_natural.group(2)
            row = (cn_num[row_str] if row_str in cn_num else int(row_str)) - 1
            count = cn_num[count_str] if count_str in cn_num else int(count_str)
            edits.append({
                "type": "merge_cells",
                "table_index": current_table,
                "start_cell": [row, 0],
                "end_cell": [row, count - 1],
                "author": None
            })
            continue
    
    return edits


def parse_style_edits(lines):
    """解析 Style Edits section"""
    edits = []
    current_style = None
    current_changes = {}

    for line in lines:
        style_match = re.match(r'^(.+?)\s*样式:', line)
        if style_match:
            if current_style and current_changes:
                edits.append({
                    "style_name": current_style,
                    "changes": current_changes.copy(),
                    "author": None
                })
            current_style = style_match.group(1)
            current_changes = {}
            continue

        if current_style is None:
            continue

        items = [x.strip() for x in line.split(',')]
        for item in items:
            for keyword, (key, value) in FORMAT_KEYWORDS.items():
                if keyword in item:
                    current_changes[key] = value
                    break
            else:
                result = parse_format_value(item)
                if result:
                    current_changes[result[0]] = result[1]

    if current_style and current_changes:
        edits.append({
            "style_name": current_style,
            "changes": current_changes,
            "author": None
        })

    return edits


def parse_global_changes(lines):
    """解析 Global Changes section"""
    replacements = []
    for line in lines:
        match = re.match(r'^将\s*"?([^"]+)"?\s*改为\s*"?([^"]+)"?', line)
        if match:
            replacements.append({
                "old": match.group(1),
                "new": match.group(2)
            })
    return replacements


def parse_md(content):
    """解析 MD 内容，返回 config dict（不含 docx 依赖操作）"""
    frontmatter, body = parse_frontmatter(content)
    sections = parse_sections(body)

    config = {
        "enable_track_revisions": frontmatter.get('track_revisions', 'true') == 'true',
        "lock_revisions": frontmatter.get('lock_revisions', 'false') == 'true',
        "revision_password": frontmatter.get('revision_password'),
        "author": frontmatter.get('author', 'Reviewer'),
        "source": frontmatter.get('source'),
        "output": frontmatter.get('output'),
        "comments": [],
        "text_modifications": [],
        "format_modifications": [],
        "table_modifications": [],
        "style_modifications": [],
        "global_replacements": [],
    }

    if 'Comments' in sections:
        config["comments"] = parse_comments(sections['Comments'])

    if 'Text Edits' in sections:
        config["text_modifications"] = parse_text_edits(sections['Text Edits'])

    if 'Format Edits' in sections:
        config["format_modifications"] = parse_format_edits(sections['Format Edits'])

    if 'Table Edits' in sections:
        config["table_modifications"] = parse_table_edits(sections['Table Edits'])

    if 'Style Edits' in sections:
        config["style_modifications"] = parse_style_edits(sections['Style Edits'])

    if 'Global Changes' in sections:
        config["global_replacements"] = parse_global_changes(sections['Global Changes'])

    return config


def test_mode():
    """Run built-in tests"""
    test_md = """---
author: Test User
source: test.docx
output: test_revised.docx
---

# Comments

Para 5: 方法论
此处的实验设计需要更多细节
> 选中范围: 第10-30字符
> 缩写: T

# Text Edits

Para 8: 标题
将 "novel" 改为 "improved"

Para 12: 插入
在开头插入: Updated: 
在末尾插入: (revised)

Para 15: 删除
删除: "old text"

Para 18: 删除有歧义
删除: "the" (第10-13字符)

# Format Edits

Para 1: 标题
居中对齐, 字号16pt, 加粗

Para 5-10: 正文
两端对齐, 行距1.5倍, 段前6pt, 左缩进12pt

# Table Edits

表格0:
  第2行下方加一行
  删掉第5行
  合并第一行的A-C列
  合并第二行的三个格子

# Style Edits

Heading1 样式:
  字号14pt, 加粗

Normal 样式:
  行距1.5倍, 段前6pt

# Global Changes

将 "significant" 改为 "statistically significant"
"""
    config = parse_md(test_md)

    errors = []
    if config['author'] != 'Test User':
        errors.append(f"Author mismatch: {config['author']}")
    if len(config['comments']) != 1:
        errors.append(f"Comments count: {len(config['comments'])}")
    if len(config['text_modifications']) != 5:
        errors.append(f"Text edits count: {len(config['text_modifications'])}")
    if len(config['format_modifications']) != 2:
        errors.append(f"Format edits count: {len(config['format_modifications'])}")
    if len(config['table_modifications']) != 4:
        errors.append(f"Table edits count: {len(config['table_modifications'])}")
    if len(config['style_modifications']) != 2:
        errors.append(f"Style edits count: {len(config['style_modifications'])}")
    if len(config['global_replacements']) != 1:
        errors.append(f"Global replacements count: {len(config['global_replacements'])}")

    if config['text_modifications'][0]['type'] != 'replace':
        errors.append(f"First edit type: {config['text_modifications'][0]['type']}")
    if config['text_modifications'][0]['old_text'] != 'novel':
        errors.append(f"First edit old_text: {config['text_modifications'][0]['old_text']}")

    c = config['comments'][0]
    if c['paragraph_index'] != 5:
        errors.append(f"Comment para_index: {c['paragraph_index']}")
    if c['start_pos'] != 10 or c['end_pos'] != 30:
        errors.append(f"Comment range: {c['start_pos']}-{c['end_pos']}")
    if c['initials'] != 'T':
        errors.append(f"Comment initials: {c['initials']}")

    f1 = config['format_modifications'][0]
    if f1['changes'].get('alignment') != 'center':
        errors.append(f"Format1 alignment: {f1['changes'].get('alignment')}")
    if f1['changes'].get('bold') != True:
        errors.append(f"Format1 bold: {f1['changes'].get('bold')}")

    t0 = config['table_modifications'][0]
    if t0['type'] != 'insert_row':
        errors.append(f"Table edit 0 type: {t0['type']}")
    if t0['row_index'] != 2:
        errors.append(f"Table edit 0 row: {t0['row_index']}")

    s0 = config['style_modifications'][0]
    if s0['style_name'] != 'Heading1':
        errors.append(f"Style 0 name: {s0['style_name']}")
    if s0['changes'].get('font_size') != 14:
        errors.append(f"Style 0 font_size: {s0['changes'].get('font_size')}")

    g = config['global_replacements'][0]
    if g['old'] != 'significant':
        errors.append(f"Global old: {g['old']}")
    if g['new'] != 'statistically significant':
        errors.append(f"Global new: {g['new']}")

    if errors:
        print("TEST FAILED:")
        for e in errors:
            print(f"  - {e}")
        return False
    else:
        print("TEST PASSED: All checks passed")
        print(f"  Comments: {len(config['comments'])}")
        print(f"  Text edits: {len(config['text_modifications'])}")
        print(f"  Format edits: {len(config['format_modifications'])}")
        print(f"  Table edits: {len(config['table_modifications'])}")
        print(f"  Style edits: {len(config['style_modifications'])}")
        print(f"  Global replacements: {len(config['global_replacements'])}")
        return True


def main():
    parser = argparse.ArgumentParser(description='将 MD 变更文件转换为 docx_revision JSON 配置')
    parser.add_argument('md_file', nargs='?', help='Markdown input file')
    parser.add_argument('json_file', nargs='?', help='JSON output file')
    parser.add_argument('--test', action='store_true', help='Run built-in tests')
    args = parser.parse_args()

    if args.test:
        success = test_mode()
        sys.exit(0 if success else 1)

    if not args.md_file or not args.json_file:
        parser.error('md_file and json_file are required (or use --test)')

    from docx import Document

    md_path = Path(args.md_file)
    json_path = Path(args.json_file)

    content = md_path.read_text(encoding='utf-8')
    config = parse_md(content)

    if config['global_replacements']:
        doc = Document(config["source"])
        for repl in config['global_replacements']:
            for i, para in enumerate(doc.paragraphs):
                if repl["old"] in para.text:
                    config["text_modifications"].append({
                        "type": "replace",
                        "paragraph_index": i,
                        "old_text": repl["old"],
                        "new_text": repl["new"],
                        "author": None
                    })
                    break

    if config["text_modifications"]:
        doc = Document(config["source"])
        errors = []
        for edit in config["text_modifications"]:
            error = check_ambiguity(edit, doc)
            if error:
                errors.append(error)

        if errors:
            print("\n".join(errors))
            sys.exit(1)

    config.pop("global_replacements", None)

    json_path.write_text(json.dumps(config, indent=2, ensure_ascii=False), encoding='utf-8')
    print(f"✅ 已生成: {json_path}")


if __name__ == '__main__':
    main()

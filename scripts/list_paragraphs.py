#!/usr/bin/env python3
"""列出 docx 文件的段落结构，用于确定 Para N 索引"""
import sys
from docx import Document

def main():
    if len(sys.argv) < 2:
        print("用法: python list_paragraphs.py <file.docx>")
        sys.exit(1)
    
    doc = Document(sys.argv[1])
    print(f"段落列表 (共{len(doc.paragraphs)}段):\n")
    for i, para in enumerate(doc.paragraphs):
        text = para.text[:80].replace('\n', ' ').strip()
        if text:
            print(f"{i}: {text}")
        else:
            print(f"{i}: (空段落)")

if __name__ == '__main__':
    main()

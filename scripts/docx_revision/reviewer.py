import copy
import hashlib
import base64
import json
import os
from datetime import datetime
from typing import Dict, Optional, Tuple, Any
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import copy


class ComprehensiveDocxReviewer:
    def __init__(self, docx_path: str, json_config_path: Optional[str] = None):
        self.docx_path = docx_path
        self.document = Document(docx_path)
        self.json_config = None
        self.comment_id_counter = 0
        self.revision_id_counter = 0

        if json_config_path and os.path.exists(json_config_path):
            with open(json_config_path, 'r', encoding='utf-8') as f:
                self.json_config = json.load(f)

        self._init_counters()

    def _init_counters(self):
        comments_part = self._get_comments_part(create_if_missing=False)
        if comments_part:
            comments = comments_part.element
            existing_ids = []
            for comment in comments.findall(qn('w:comment')):
                existing_ids.append(int(comment.get(qn('w:id'))))
            self.comment_id_counter = max(existing_ids, default=-1) + 1

        self.revision_id_counter = self._get_max_revision_id() + 1

    def _get_max_revision_id(self) -> int:
        max_id = 0
        
        for paragraph in self.document.paragraphs:
            for run in paragraph.runs:
                r = run._r
                for elem in r.xpath('.//w:ins | .//w:del | .//w:moveFrom | .//w:moveTo'):
                    id_val = elem.get(qn('w:id'))
                    if id_val and id_val.isdigit():
                        max_id = max(max_id, int(id_val))
        
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            r = run._r
                            for elem in r.xpath('.//w:ins | .//w:del | .//w:moveFrom | .//w:moveTo'):
                                id_val = elem.get(qn('w:id'))
                                if id_val and id_val.isdigit():
                                    max_id = max(max_id, int(id_val))
        
        return max_id

    def enable_track_revisions(self, enable: bool = True):
        settings = self.document.settings.element

        track_revisions = settings.find(qn('w:trackRevisions'))
        if track_revisions is None:
            track_revisions = OxmlElement('w:trackRevisions')
            settings.append(track_revisions)

        track_revisions.set(qn('w:val'), '1' if enable else '0')

        status = "启用" if enable else "禁用"
        print(f"✅ 修订跟踪已{status}")

    def lock_revisions(self, lock: bool = True, password: Optional[str] = None):
        settings = self.document.settings.element

        lock_revisions = settings.find(qn('w:lockRevisions'))
        if lock_revisions is None:
            lock_revisions = OxmlElement('w:lockRevisions')
            settings.append(lock_revisions)

        lock_revisions.set(qn('w:val'), '1' if lock else '0')

        if password and lock:
            password_hash = hashlib.sha1(password.encode('utf-8')).digest()
            encoded_hash = base64.b64encode(password_hash).decode('ascii')
            lock_revisions.set(qn('w:hash'), encoded_hash)

        status = "锁定" if lock else "解锁"
        print(f"✅ 修订已{status}")

    def insert_text_with_tracking(self,
                                  paragraph_index: int,
                                  text: str,
                                  position: Optional[int] = None,
                                  author: str = "Reviewer",
                                  color: Tuple[int, int, int] = (0, 0, 255)) -> int:
        if paragraph_index >= len(self.document.paragraphs):
            raise ValueError(f"段落索引 {paragraph_index} 超出范围")

        paragraph = self.document.paragraphs[paragraph_index]
        revision_id = self._get_next_revision_id()

        ins_element = OxmlElement('w:ins')
        ins_element.set(qn('w:id'), str(revision_id))
        ins_element.set(qn('w:author'), author)
        ins_element.set(qn('w:date'), self._get_current_time())

        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        color_elem = OxmlElement('w:color')
        color_elem.set(qn('w:val'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
        rPr.append(color_elem)

        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        rPr.append(underline)

        r.append(rPr)

        t = OxmlElement('w:t')
        t.text = text
        r.append(t)

        ins_element.append(r)

        if position is not None and position > 0:
            full_text = paragraph.text
            if position <= len(full_text):
                paragraph._p.clear()
                
                before_text = full_text[:position]
                after_text = full_text[position:]
                
                if before_text:
                    before_run = OxmlElement('w:r')
                    before_t = OxmlElement('w:t')
                    before_t.text = before_text
                    before_run.append(before_t)
                    paragraph._p.append(before_run)
                
                paragraph._p.append(ins_element)
                
                if after_text:
                    after_run = OxmlElement('w:r')
                    after_t = OxmlElement('w:t')
                    after_t.text = after_text
                    after_run.append(after_t)
                    paragraph._p.append(after_run)
            else:
                paragraph._p.append(ins_element)
        else:
            paragraph._p.append(ins_element)

        print(f"✅ 已插入文本（修订ID: {revision_id}）: '{text[:20]}...'")
        return revision_id

    def delete_text_with_tracking(self,
                                  paragraph_index: int,
                                  start_pos: int,
                                  end_pos: int,
                                  author: str = "Reviewer",
                                  color: Tuple[int, int, int] = (255, 0, 0)) -> int:
        if paragraph_index >= len(self.document.paragraphs):
            raise ValueError(f"段落索引 {paragraph_index} 超出范围")

        paragraph = self.document.paragraphs[paragraph_index]
        full_text = paragraph.text

        if start_pos < 0 or end_pos > len(full_text) or start_pos >= end_pos:
            raise ValueError(f"无效的删除范围: {start_pos}-{end_pos}")

        revision_id = self._get_next_revision_id()
        deleted_text = full_text[start_pos:end_pos]

        del_element = OxmlElement('w:del')
        del_element.set(qn('w:id'), str(revision_id))
        del_element.set(qn('w:author'), author)
        del_element.set(qn('w:date'), self._get_current_time())

        before_text = full_text[:start_pos]
        after_text = full_text[end_pos:]

        for run in paragraph.runs:
            run._r.getparent().remove(run._r)

        if before_text:
            paragraph.add_run(before_text)

        del_run = OxmlElement('w:r')
        del_rPr = OxmlElement('w:rPr')

        strike = OxmlElement('w:strike')
        strike.set(qn('w:val'), '1')
        del_rPr.append(strike)

        color_elem = OxmlElement('w:color')
        color_elem.set(qn('w:val'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
        del_rPr.append(color_elem)

        del_run.append(del_rPr)

        del_t = OxmlElement('w:delText')
        del_t.text = deleted_text
        del_run.append(del_t)

        del_element.append(del_run)
        paragraph._p.append(del_element)

        if after_text:
            paragraph.add_run(after_text)

        print(f"✅ 已删除文本（修订ID: {revision_id}）: '{deleted_text[:20]}...'")
        return revision_id

    def replace_text_with_tracking(self,
                                   paragraph_index: int,
                                   old_text: str,
                                   new_text: str,
                                   author: str = "Reviewer",
                                   delete_color: Tuple[int, int, int] = (255, 0, 0),
                                   insert_color: Tuple[int, int, int] = (0, 0, 255)) -> Tuple[int, int]:
        if paragraph_index >= len(self.document.paragraphs):
            raise ValueError(f"段落索引 {paragraph_index} 超出范围")

        paragraph = self.document.paragraphs[paragraph_index]
        full_text = paragraph.text

        if old_text not in full_text:
            raise ValueError(f"未找到文本: {old_text}")

        start_pos = full_text.index(old_text)
        end_pos = start_pos + len(old_text)

        del_id = self.delete_text_with_tracking(
            paragraph_index, start_pos, end_pos, author, delete_color
        )

        ins_id = self.insert_text_with_tracking(
            paragraph_index, new_text, None, author, insert_color
        )

        print(f"✅ 已替换文本: '{old_text[:20]}...' → '{new_text[:20]}...'")
        return del_id, ins_id

    def change_paragraph_format_with_tracking(self,
                                             paragraph_index: int,
                                             format_changes: Dict[str, Any],
                                             author: str = "Reviewer") -> int:
        if paragraph_index >= len(self.document.paragraphs):
            raise ValueError(f"段落索引 {paragraph_index} 超出范围")

        paragraph = self.document.paragraphs[paragraph_index]
        revision_id = self._get_next_revision_id()

        p = paragraph._p
        pPr = p.get_or_add_pPr()

        pPr_change = OxmlElement('w:pPrChange')
        pPr_change.set(qn('w:id'), str(revision_id))
        pPr_change.set(qn('w:author'), author)
        pPr_change.set(qn('w:date'), self._get_current_time())

        old_pPr = OxmlElement('w:pPr')
        for child in list(pPr):
            if child.tag != qn('w:pPrChange'):
                old_pPr.append(copy.deepcopy(child))
        pPr_change.append(old_pPr)
        pPr.append(pPr_change)

        for prop, value in format_changes.items():
            if prop == 'alignment':
                jc = pPr.find(qn('w:jc'))
                if jc is None:
                    jc = OxmlElement('w:jc')
                    pPr.insert(len(pPr) - 1, jc)
                jc.set(qn('w:val'), value)
            elif prop == 'line_spacing':
                spacing = pPr.find(qn('w:spacing'))
                if spacing is None:
                    spacing = OxmlElement('w:spacing')
                    pPr.insert(len(pPr) - 1, spacing)
                spacing.set(qn('w:line'), str(int(value * 240)))
                spacing.set(qn('w:lineRule'), 'auto')
            elif prop == 'space_before':
                spacing = pPr.find(qn('w:spacing'))
                if spacing is None:
                    spacing = OxmlElement('w:spacing')
                    pPr.insert(len(pPr) - 1, spacing)
                spacing.set(qn('w:before'), str(int(value * 20)))
            elif prop == 'space_after':
                spacing = pPr.find(qn('w:spacing'))
                if spacing is None:
                    spacing = OxmlElement('w:spacing')
                    pPr.insert(len(pPr) - 1, spacing)
                spacing.set(qn('w:after'), str(int(value * 20)))
            elif prop == 'indent_left':
                ind = pPr.find(qn('w:ind'))
                if ind is None:
                    ind = OxmlElement('w:ind')
                    pPr.insert(len(pPr) - 1, ind)
                ind.set(qn('w:left'), str(int(value * 20)))
            elif prop == 'indent_right':
                ind = pPr.find(qn('w:ind'))
                if ind is None:
                    ind = OxmlElement('w:ind')
                    pPr.insert(len(pPr) - 1, ind)
                ind.set(qn('w:right'), str(int(value * 20)))

        changes_desc = ', '.join([f'{k}={v}' for k, v in format_changes.items()])
        print(f"✅ 已更改段落格式（修订ID: {revision_id}）: {changes_desc}")
        return revision_id

    def insert_table_row_with_tracking(self,
                                       table_index: int,
                                       row_index: int,
                                       author: str = "Reviewer") -> int:
        tables = self.document.tables
        if table_index >= len(tables):
            raise ValueError(f"表格索引 {table_index} 超出范围")

        table = tables[table_index]
        revision_id = self._get_next_revision_id()

        tbl = table._tbl
        tr_elements = tbl.findall(qn('w:tr'))

        new_tr = copy.deepcopy(tr_elements[-1])
        for tc in new_tr.findall(qn('w:tc')):
            for p in tc.findall(qn('w:p')):
                tc.remove(p)
            new_p = OxmlElement('w:p')
            tc.append(new_p)

        ins_element = OxmlElement('w:ins')
        ins_element.set(qn('w:id'), str(revision_id))
        ins_element.set(qn('w:author'), author)
        ins_element.set(qn('w:date'), self._get_current_time())
        new_tr.insert(0, ins_element)

        tbl.insert(row_index, new_tr)

        print(f"✅ 已插入表格行（修订ID: {revision_id}），位置: {row_index}")
        return revision_id

    def delete_table_row_with_tracking(self,
                                      table_index: int,
                                      row_index: int,
                                      author: str = "Reviewer") -> int:
        tables = self.document.tables
        if table_index >= len(tables):
            raise ValueError(f"表格索引 {table_index} 超出范围")

        table = tables[table_index]
        if row_index >= len(table.rows):
            raise ValueError(f"行索引 {row_index} 超出范围")

        revision_id = self._get_next_revision_id()
        date_str = self._get_current_time()

        tbl = table._tbl
        tr = table.rows[row_index]._tr

        for tc in tr.findall(qn('w:tc')):
            for p in tc.findall(qn('w:p')):
                runs = p.findall(qn('w:r'))
                for r in runs:
                    t_elem = r.find(qn('w:t'))
                    if t_elem is not None and t_elem.text:
                        del_elem = OxmlElement('w:del')
                        del_elem.set(qn('w:id'), str(revision_id))
                        del_elem.set(qn('w:author'), author)
                        del_elem.set(qn('w:date'), date_str)
                        del_r = copy.deepcopy(r)
                        del_t = del_r.find(qn('w:t'))
                        if del_t is not None:
                            del_t.tag = qn('w:delText')
                        p.remove(r)
                        del_elem.append(del_r)
                        p.append(del_elem)

        print(f"✅ 已删除表格行（修订ID: {revision_id}），位置: {row_index}")
        return revision_id

    def merge_cells_with_tracking(self,
                                 table_index: int,
                                 start_cell: Tuple[int, int],
                                 end_cell: Tuple[int, int],
                                 author: str = "Reviewer") -> int:
        tables = self.document.tables
        if table_index >= len(tables):
            raise ValueError(f"表格索引 {table_index} 超出范围")

        table = tables[table_index]
        revision_id = self._get_next_revision_id()

        cell_merge = OxmlElement('w:cellMerge')
        cell_merge.set(qn('w:id'), str(revision_id))
        cell_merge.set(qn('w:author'), author)
        cell_merge.set(qn('w:date'), self._get_current_time())

        start_row, start_col = start_cell
        end_row, end_col = end_cell

        start_cell_obj = table.cell(start_row, start_col)

        if start_row == end_row:
            for col in range(start_col + 1, end_col + 1):
                start_cell_obj.merge(table.cell(end_row, col))
        else:
            for row in range(start_row, end_row + 1):
                for col in range(start_col, end_col + 1):
                    if row == start_row and col == start_col:
                        continue
                    start_cell_obj.merge(table.cell(row, col))

        tc = start_cell_obj._tc
        tcPr = tc.get_or_add_tcPr()
        tcPr.append(cell_merge)

        print(f"✅ 已合并单元格（修订ID: {revision_id}）: ({start_row},{start_col}) 到 ({end_row},{end_col})")
        return revision_id

    def add_comment(self,
                    paragraph_index: int,
                    comment_text: str,
                    start_pos: Optional[int] = None,
                    end_pos: Optional[int] = None,
                    author: str = "Reviewer",
                    initials: str = "RV") -> int:
        if paragraph_index >= len(self.document.paragraphs):
            raise ValueError(f"段落索引 {paragraph_index} 超出范围")

        paragraph = self.document.paragraphs[paragraph_index]
        comment_id = self._get_next_comment_id()

        full_text = paragraph.text
        if start_pos is None:
            start_pos = 0
        if end_pos is None:
            end_pos = len(full_text)

        if start_pos < 0 or end_pos > len(full_text) or start_pos >= end_pos:
            raise ValueError(f"无效的批注范围: {start_pos}-{end_pos}")

        comment_range_start = OxmlElement('w:commentRangeStart')
        comment_range_start.set(qn('w:id'), str(comment_id))

        comment_range_end = OxmlElement('w:commentRangeEnd')
        comment_range_end.set(qn('w:id'), str(comment_id))

        self._insert_comment_range_marks(paragraph, comment_range_start, comment_range_end, start_pos, end_pos)

        comment_ref = OxmlElement('w:r')
        comment_ref_pr = OxmlElement('w:rPr')
        comment_ref_id = OxmlElement('w:commentReference')
        comment_ref_id.set(qn('w:id'), str(comment_id))
        comment_ref_pr.append(comment_ref_id)
        comment_ref.append(comment_ref_pr)
        paragraph._p.append(comment_ref)

        self._add_comment_content(comment_id, comment_text, author, initials)

        print(f"✅ 已添加批注（ID: {comment_id}）: '{comment_text[:30]}...'")
        return comment_id

    def modify_style_with_tracking(self,
                                   style_name: str,
                                   style_type: str = "paragraph",
                                   format_changes: Dict[str, Any] = None,
                                   author: str = "Reviewer") -> int:
        revision_id = self._get_next_revision_id()

        valid_style_types = {"paragraph", "character", "table", "list"}
        if style_type not in valid_style_types:
            raise ValueError(f"不支持的样式类型: {style_type}")

        style = self.document.styles[style_name]

        style_change = OxmlElement('w:styleChange')
        style_change.set(qn('w:id'), str(revision_id))
        style_change.set(qn('w:author'), author)
        style_change.set(qn('w:date'), self._get_current_time())

        if format_changes:
            if style_type == "paragraph":
                pPr = style.element.get_or_add_pPr()
                for prop, value in format_changes.items():
                    if prop == 'alignment':
                        jc = OxmlElement('w:jc')
                        jc.set(qn('w:val'), value)
                        pPr.append(jc)
                    elif prop == 'space_before':
                        spacing = pPr.get_or_add_spacing()
                        spacing.set(qn('w:before'), str(int(value * 20)))
                    elif prop == 'space_after':
                        spacing = pPr.get_or_add_spacing()
                        spacing.set(qn('w:after'), str(int(value * 20)))

            elif style_type == "character":
                rPr = style.element.get_or_add_rPr()
                for prop, value in format_changes.items():
                    if prop == 'bold' and value:
                        b = OxmlElement('w:b')
                        b.set(qn('w:val'), '1')
                        rPr.append(b)
                    elif prop == 'italic' and value:
                        i = OxmlElement('w:i')
                        i.set(qn('w:val'), '1')
                        rPr.append(i)
                    elif prop == 'font_size':
                        sz = OxmlElement('w:sz')
                        sz.set(qn('w:val'), str(int(value * 2)))
                        rPr.append(sz)
                    elif prop == 'font_name':
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:ascii'), value)
                        rFonts.set(qn('w:hAnsi'), value)
                        rPr.append(rFonts)

        style.element.append(style_change)

        changes_desc = ', '.join([f'{k}={v}' for k, v in (format_changes or {}).items()])
        print(f"✅ 已修改样式 '{style_name}'（修订ID: {revision_id}）: {changes_desc}")
        return revision_id

    def apply_json_config(self, config: Optional[Dict] = None):
        if config:
            self.json_config = config
        elif not self.json_config:
            raise ValueError("未提供JSON配置")

        cfg = self.json_config

        if cfg.get('enable_track_revisions', True):
            self.enable_track_revisions(True)

        if cfg.get('lock_revisions', False):
            self.lock_revisions(True, cfg.get('revision_password'))

        for comment in cfg.get('comments', []):
            self.add_comment(
                paragraph_index=comment['paragraph_index'],
                comment_text=comment['text'],
                start_pos=comment.get('start_pos'),
                end_pos=comment.get('end_pos'),
                author=comment.get('author', 'Reviewer'),
                initials=comment.get('initials', 'RV')
            )

        for mod in cfg.get('text_modifications', []):
            mod_type = mod.get('type', 'replace')

            if mod_type == 'insert':
                self.insert_text_with_tracking(
                    paragraph_index=mod['paragraph_index'],
                    text=mod['text'],
                    position=mod.get('position'),
                    author=mod.get('author', 'Reviewer'),
                    color=tuple(mod.get('color', [0, 0, 255]))
                )
            elif mod_type == 'delete':
                start_pos = mod.get('start_pos')
                end_pos = mod.get('end_pos')
                
                if start_pos is None or end_pos is None:
                    paragraph = self.document.paragraphs[mod['paragraph_index']]
                    text_to_delete = mod['text']
                    pos = paragraph.text.find(text_to_delete)
                    if pos == -1:
                        print(f"⚠️ 未找到文本: '{text_to_delete[:30]}...' 在段落 {mod['paragraph_index']}")
                        continue
                    start_pos = pos
                    end_pos = pos + len(text_to_delete)
                
                self.delete_text_with_tracking(
                    paragraph_index=mod['paragraph_index'],
                    start_pos=start_pos,
                    end_pos=end_pos,
                    author=mod.get('author', 'Reviewer'),
                    color=tuple(mod.get('color', [255, 0, 0]))
                )
            elif mod_type == 'replace':
                self.replace_text_with_tracking(
                    paragraph_index=mod['paragraph_index'],
                    old_text=mod['old_text'],
                    new_text=mod['new_text'],
                    author=mod.get('author', 'Reviewer'),
                    delete_color=tuple(mod.get('delete_color', [255, 0, 0])),
                    insert_color=tuple(mod.get('insert_color', [0, 0, 255]))
                )
            # ponytail: removed 'move' type - stubs had no effect

        for fmt in cfg.get('format_modifications', []):
            if fmt.get('scope') == 'paragraph':
                self.change_paragraph_format_with_tracking(
                    paragraph_index=fmt['paragraph_index'],
                    format_changes=fmt['changes'],
                    author=fmt.get('author', 'Reviewer')
                )
            # ponytail: removed 'text' scope - rPrChange was never appended to XML

        for tbl_mod in cfg.get('table_modifications', []):
            if tbl_mod['type'] == 'insert_row':
                self.insert_table_row_with_tracking(
                    table_index=tbl_mod['table_index'],
                    row_index=tbl_mod['row_index'],
                    author=tbl_mod.get('author', 'Reviewer')
                )
            elif tbl_mod['type'] == 'delete_row':
                self.delete_table_row_with_tracking(
                    table_index=tbl_mod['table_index'],
                    row_index=tbl_mod['row_index'],
                    author=tbl_mod.get('author', 'Reviewer')
                )
            elif tbl_mod['type'] == 'merge_cells':
                self.merge_cells_with_tracking(
                    table_index=tbl_mod['table_index'],
                    start_cell=tuple(tbl_mod['start_cell']),
                    end_cell=tuple(tbl_mod['end_cell']),
                    author=tbl_mod.get('author', 'Reviewer')
                )

        for style_mod in cfg.get('style_modifications', []):
            self.modify_style_with_tracking(
                style_name=style_mod['style_name'],
                style_type=style_mod.get('style_type', 'paragraph'),
                format_changes=style_mod.get('changes', {}),
                author=style_mod.get('author', 'Reviewer')
            )

        print("✅ JSON配置应用完成")

    def _get_next_revision_id(self) -> int:
        current_id = self.revision_id_counter
        self.revision_id_counter += 1
        return current_id

    def _get_next_comment_id(self) -> int:
        current_id = self.comment_id_counter
        self.comment_id_counter += 1
        return current_id

    def _get_current_time(self) -> str:
        return datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ')

    def _get_comments_part(self, create_if_missing: bool = True):
        try:
            return self.document.part.part_related_by(RT.COMMENTS)
        except KeyError:
            if not create_if_missing:
                return None

            from docx.opc.part import XmlPart
            from docx.opc.packuri import PackURI
            from docx.oxml.parser import parse_xml

            comments_element = parse_xml(
                '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
            comments_part = XmlPart(
                PackURI('/word/comments.xml'),
                'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml',
                comments_element,
                self.document.part.package
            )
            self.document.part.relate_to(comments_part, RT.COMMENTS)

            RT_COMMENTS_EXT = 'http://schemas.microsoft.com/office/2011/relationships/commentsExtended'

            comments_ext_element = parse_xml(
                '<w15:commentsEx xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"/>'
            )
            comments_ext_part = XmlPart(
                PackURI('/word/commentsExtended.xml'),
                'application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtended+xml',
                comments_ext_element,
                self.document.part.package
            )
            self.document.part.relate_to(comments_ext_part, RT_COMMENTS_EXT)

            return comments_part

    def _add_comment_content(self, comment_id: int, text: str, author: str, initials: str):
        comments_part = self._get_comments_part()
        comments = comments_part.element

        comment = OxmlElement('w:comment')
        comment.set(qn('w:id'), str(comment_id))
        comment.set(qn('w:author'), author)
        comment.set(qn('w:initials'), initials)
        comment.set(qn('w:date'), self._get_current_time())

        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        p.append(r)
        comment.append(p)

        comments.append(comment)
        comments_part._element = comments

        para_id = f'{0x5E23DD26 + comment_id:08X}'

        RT_COMMENTS_EXT = 'http://schemas.microsoft.com/office/2011/relationships/commentsExtended'

        from docx.oxml.parser import parse_xml

        comments_ext_part = self.document.part.part_related_by(RT_COMMENTS_EXT)
        comments_ext = comments_ext_part.element

        comment_ex = parse_xml(
            f'<w15:commentEx w15:paraId="{para_id}" w15:done="0"'
            ' xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"/>'
        )
        comments_ext.append(comment_ex)
        comments_ext_part._element = comments_ext

    def _insert_comment_range_marks(self, paragraph: 'Paragraph', start_mark: 'BaseOxmlElement',
                                   end_mark: 'BaseOxmlElement', start_pos: int, end_pos: int):
        p = paragraph._p

        first_run = p.find(qn('w:r'))
        if first_run is not None:
            p.insert(list(p).index(first_run), start_mark)
        else:
            p.insert(0, start_mark)

        p.append(end_mark)

    def save(self, output_path: Optional[str] = None) -> str:
        if output_path is None:
            base_name = os.path.splitext(self.docx_path)[0]
            output_path = f"{base_name}_reviewed.docx"

        self.document.save(output_path)
        print(f"✅ 文档已保存到: {output_path}")
        return output_path

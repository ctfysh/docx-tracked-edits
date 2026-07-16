from docx import Document
from .reviewer import ComprehensiveDocxReviewer


def main():
    doc = Document()
    doc.add_heading('项目进展报告', 0)
    doc.add_paragraph('日期：2024年1月15日')
    doc.add_paragraph('这是第一个需要修改的段落。包含一些需要删除的文本。')
    doc.add_paragraph('第二个段落，包含一些需要替换的内容，比如约50%这样的数据。')
    doc.add_paragraph('第三个段落，用于演示段落格式变更。')
    doc.add_paragraph('第四个段落，用于演示表格操作。')

    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    for row in range(3):
        for col in range(3):
            table.cell(row, col).text = f'单元格 ({row},{col})'

    doc.save('sample_input.docx')

    reviewer = ComprehensiveDocxReviewer('sample_input.docx')
    reviewer.enable_track_revisions()

    reviewer.add_comment(0, "标题需要更具体，建议包含项目名称和版本号。", author="项目经理", initials="PM")
    reviewer.add_comment(3, "准确率数据需要注明测试条件和数据集。", author="技术总监", initials="TD")

    reviewer.insert_text_with_tracking(1, "[紧急] ", position=0, author="主管", color=(255, 0, 0))
    reviewer.delete_text_with_tracking(2, 10, 25, author="编辑", color=(255, 0, 0))
    reviewer.replace_text_with_tracking(3, "约50%", "精确值为52.3%", author="数据分析师",
                                       delete_color=(255, 0, 0), insert_color=(0, 128, 0))

    reviewer.change_paragraph_format_with_tracking(2, {
        'alignment': 'center', 'line_spacing': 1.5,
        'space_before': 12, 'space_after': 12,
        'indent_left': 36, 'indent_right': 24
    }, author="排版专家")

    reviewer.insert_table_row_with_tracking(0, 2, author="数据录入员")
    reviewer.delete_table_row_with_tracking(0, 2, author="数据审核员")
    reviewer.merge_cells_with_tracking(0, (1, 0), (1, 2), author="表格设计师")

    reviewer.modify_style_with_tracking("Normal", "paragraph", {'space_before': 6, 'space_after': 6})
    reviewer.modify_style_with_tracking("Emphasis", "character", {
        'bold': True, 'italic': True, 'font_size': 12,
        'color': [0, 0, 255], 'font_name': 'Arial'
    })

    reviewer.save('comprehensive_output.docx')


if __name__ == '__main__':
    main()

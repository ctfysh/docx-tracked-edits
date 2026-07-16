#!/usr/bin/env python3
"""Tests for docx_revision package"""

import os
import sys
import json
import tempfile
import shutil
from pathlib import Path

from docx import Document
from .reviewer import ComprehensiveDocxReviewer


class TestComprehensiveDocxReviewer:

    def __init__(self):
        self.test_dir = tempfile.mkdtemp()
        self.sample_docx = os.path.join(self.test_dir, 'test_sample.docx')
        self.output_docx = os.path.join(self.test_dir, 'test_output.docx')
        self.config_json = os.path.join(self.test_dir, 'test_config.json')

        self._create_sample_document()
        self.reviewer = None

    def _create_sample_document(self):
        doc = Document()

        doc.add_heading('测试文档标题', 0)
        doc.add_paragraph('第一个段落，用于测试各种操作。')
        doc.add_paragraph('第二个段落，包含一些需要修改的文本内容。')
        doc.add_paragraph('第三个段落，用于演示格式变更。')
        doc.add_paragraph('第四个段落，用于移动操作测试。')
        doc.add_paragraph('第五个段落，移动操作的目标位置。')

        table = doc.add_table(rows=3, cols=3)
        table.style = 'Table Grid'
        for row in range(3):
            for col in range(3):
                cell = table.cell(row, col)
                cell.text = f'单元格 ({row},{col})'

        doc.save(self.sample_docx)
        print(f"✅ 测试文档已创建: {self.sample_docx}")

    def cleanup(self):
        if os.path.exists(self.test_dir):
            shutil.rmtree(self.test_dir)
            print(f"✅ 测试目录已清理: {self.test_dir}")

    def test_initialization(self):
        print("\n🔧 测试初始化...")
        self.reviewer = ComprehensiveDocxReviewer(self.sample_docx)

        assert self.reviewer.docx_path == self.sample_docx
        assert self.reviewer.document is not None
        assert self.reviewer.comment_id_counter >= 0
        assert self.reviewer.revision_id_counter >= 0

        print("✅ 初始化测试通过")
        return True

    def test_enable_track_revisions(self):
        print("\n🔧 测试启用/禁用修订跟踪...")

        self.reviewer.enable_track_revisions(True)

        settings = self.reviewer.document.settings.element
        track_revisions = settings.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}trackRevisions')
        assert track_revisions is not None
        assert track_revisions.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') == '1'

        self.reviewer.enable_track_revisions(False)
        track_revisions = settings.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}trackRevisions')
        assert track_revisions.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') == '0'

        print("✅ 启用/禁用修订跟踪测试通过")
        return True

    def test_lock_revisions(self):
        print("\n🔧 测试锁定/解锁修订...")

        self.reviewer.lock_revisions(True, "test_password")

        settings = self.reviewer.document.settings.element
        lock_revisions = settings.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lockRevisions')
        assert lock_revisions is not None
        assert lock_revisions.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') == '1'

        self.reviewer.lock_revisions(False)
        lock_revisions = settings.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lockRevisions')
        assert lock_revisions.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') == '0'

        print("✅ 锁定/解锁修订测试通过")
        return True

    def test_insert_text_with_tracking(self):
        print("\n🔧 测试插入文本（修订）...")

        revision_id = self.reviewer.insert_text_with_tracking(
            paragraph_index=1,
            text="[测试插入] ",
            position=0,
            author="测试员"
        )

        assert revision_id >= 0

        paragraph = self.reviewer.document.paragraphs[1]
        xml = paragraph._p.xml
        print(f"DEBUG: Paragraph XML: {xml[:500]}...")

        if "<w:ins" not in xml:
            print("DEBUG: <w:ins> not found in XML")
            return False
        if "[测试插入]" not in xml:
            print("DEBUG: [测试插入] not found in XML")
            return False

        print(f"✅ 插入文本测试通过 (修订ID: {revision_id})")
        return True

    def test_delete_text_with_tracking(self):
        print("\n🔧 测试删除文本（修订）...")

        paragraph = self.reviewer.document.paragraphs[2]
        original_text = paragraph.text
        original_length = len(original_text)

        revision_id = self.reviewer.delete_text_with_tracking(
            paragraph_index=2,
            start_pos=5,
            end_pos=15,
            author="测试员"
        )

        assert revision_id >= 0

        paragraph = self.reviewer.document.paragraphs[2]
        assert paragraph is not None

        print(f"✅ 删除文本测试通过 (修订ID: {revision_id})")
        return True

    def test_replace_text_with_tracking(self):
        print("\n🔧 测试替换文本（修订）...")

        del_id, ins_id = self.reviewer.replace_text_with_tracking(
            paragraph_index=3,
            old_text="格式",
            new_text="样式",
            author="测试员"
        )

        assert del_id >= 0
        assert ins_id >= 0

        paragraph = self.reviewer.document.paragraphs[3]
        xml = paragraph._p.xml
        assert "<w:del" in xml
        assert "<w:ins" in xml
        assert "样式" in xml

        print(f"✅ 替换文本测试通过 (删除ID: {del_id}, 插入ID: {ins_id})")
        return True

    def test_change_paragraph_format_with_tracking(self):
        print("\n🔧 测试段落格式变更（修订）...")

        revision_id = self.reviewer.change_paragraph_format_with_tracking(
            paragraph_index=2,
            format_changes={
                'alignment': 'center',
                'space_before': 12,
                'space_after': 12,
                'indent_left': 36
            },
            author="测试员"
        )

        assert revision_id >= 0

        print(f"✅ 段落格式变更测试通过 (修订ID: {revision_id})")
        return True

    def test_table_operations(self):
        print("\n🔧 测试表格操作（修订）...")

        insert_id = self.reviewer.insert_table_row_with_tracking(
            table_index=0,
            row_index=1,
            author="测试员"
        )
        assert insert_id >= 0

        delete_id = self.reviewer.delete_table_row_with_tracking(
            table_index=0,
            row_index=2,
            author="测试员"
        )
        assert delete_id >= 0

        merge_id = self.reviewer.merge_cells_with_tracking(
            table_index=0,
            start_cell=(0, 0),
            end_cell=(0, 2),
            author="测试员"
        )
        assert merge_id >= 0

        print(f"✅ 表格操作测试通过 (插入ID: {insert_id}, 删除ID: {delete_id}, 合并ID: {merge_id})")
        return True

    def test_comment_operations(self):
        print("\n🔧 测试批注操作...")

        comment_id = self.reviewer.add_comment(
            paragraph_index=1,
            comment_text="这是一个测试批注。",
            start_pos=0,
            end_pos=10,
            author="测试员",
            initials="T"
        )
        assert comment_id >= 0

        comments_part = self.reviewer._get_comments_part(create_if_missing=False)
        if comments_part:
            comments_xml = comments_part.element.xml
            assert f'w:id="{comment_id}"' in comments_xml

        print(f"✅ 批注操作测试通过 (批注ID: {comment_id})")
        return True

    def test_modify_style_with_tracking(self):
        print("\n🔧 测试样式修改（修订）...")

        revision_id = self.reviewer.modify_style_with_tracking(
            style_name="Normal",
            style_type="paragraph",
            format_changes={
                'space_before': 6,
                'space_after': 6
            },
            author="测试员"
        )

        assert revision_id >= 0

        print(f"✅ 样式修改测试通过 (修订ID: {revision_id})")
        return True

    def test_apply_json_config(self):
        print("\n🔧 测试JSON配置应用...")

        config = {
            "enable_track_revisions": True,
            "lock_revisions": False,
            "revision_password": None,
            "comments": [
                {
                    "paragraph_index": 0,
                    "text": "JSON配置测试批注。",
                    "start_pos": 0,
                    "end_pos": 5,
                    "author": "配置测试员",
                    "initials": "C"
                }
            ],
            "text_modifications": [
                {
                    "type": "insert",
                    "paragraph_index": 1,
                    "text": "[配置插入] ",
                    "position": 0,
                    "author": "配置测试员",
                    "color": [255, 0, 0]
                }
            ]
        }

        with open(self.config_json, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)

        reviewer = ComprehensiveDocxReviewer(self.sample_docx, self.config_json)
        reviewer.apply_json_config()

        paragraph = reviewer.document.paragraphs[1]
        xml = paragraph._p.xml
        assert "<w:ins" in xml
        assert "[配置插入]" in xml

        print("✅ JSON配置应用测试通过")
        return True

    def test_save_document(self):
        print("\n🔧 测试保存文档...")

        output_path = self.reviewer.save(self.output_docx)

        assert os.path.exists(output_path)
        assert os.path.getsize(output_path) > 0

        print(f"✅ 保存文档测试通过: {output_path}")
        return True

    def run_all_tests(self):
        print("🚀 开始运行全面测试...")

        tests = [
            self.test_initialization,
            self.test_enable_track_revisions,
            self.test_lock_revisions,
            self.test_insert_text_with_tracking,
            self.test_delete_text_with_tracking,
            self.test_replace_text_with_tracking,
            self.test_change_paragraph_format_with_tracking,
            self.test_table_operations,
            self.test_comment_operations,
            self.test_modify_style_with_tracking,
            self.test_apply_json_config,
            self.test_save_document
        ]

        passed = 0
        failed = 0

        for test in tests:
            try:
                if test():
                    passed += 1
                else:
                    failed += 1
            except Exception as e:
                print(f"❌ 测试失败: {test.__name__} - {e}")
                failed += 1

        print(f"\n📊 测试结果: {passed} 通过, {failed} 失败")

        self.cleanup()

        return failed == 0


def main():
    tester = TestComprehensiveDocxReviewer()
    success = tester.run_all_tests()

    if success:
        print("\n🎉 所有测试通过！")
        return 0
    else:
        print("\n💥 有测试失败！")
        return 1


if __name__ == '__main__':
    sys.exit(main())

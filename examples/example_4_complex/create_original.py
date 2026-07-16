#!/usr/bin/env python3
"""生成 Example 4: 多节复杂文档的原始 docx 文档

段落结构 (表格不计入段落索引):
  Para 0:   标题 "Research Project Proposal" (Heading level 0)
  Para 1-2: 空段落
  Para 3:   项目概述 (含 "项目组")
  Para 4-7: 空段落
  Para 8:   技术方案 (含 "项目组")
  Para 9-11: 空段落
  Para 12:  预算表标题 "表1：预算分配"
  [表格: 3行x3列: 项目/金额/说明]
  Para 13-14: 空段落
  Para 15:  时间计划
  Para 16-19: 空段落
  Para 20:  项目团队 (含 "Tiger公司" 和 "项目组")
  Para 21-24: 空段落
  Para 25:  参考文献
"""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    print("需要安装 python-docx: pip install python-docx")
    sys.exit(1)


def create_original_doc(output_path: str = None):
    """生成原始复杂文档 docx"""
    if output_path is None:
        output_path = str(Path(__file__).parent / "example_4_original.docx")

    doc = Document()

    # Para 0: 标题
    doc.add_heading("Research Project Proposal", level=0)

    # Para 1-2: 空段落
    doc.add_paragraph("")  # Para 1
    doc.add_paragraph("")  # Para 2

    # Para 3: 项目概述 (含 "项目组")
    doc.add_paragraph(
        "本项目旨在开发一套智能水质监测系统，利用物联网和大数据技术"
        "实现对湖泊水质的实时监控和预警。项目组将采用先进的传感器网络，"
        "覆盖关键监测点，确保数据采集的全面性和准确性。"
    )

    # Para 4-7: 空段落
    doc.add_paragraph("")  # Para 4
    doc.add_paragraph("")  # Para 5
    doc.add_paragraph("")  # Para 6
    doc.add_paragraph("")  # Para 7

    # Para 8: 技术方案 (含 "项目组")
    doc.add_paragraph(
        "技术方案：本项目采用多层次架构设计，包括数据采集层、数据传输层、"
        "数据处理层和应用展示层。项目组将基于云计算平台构建数据处理管道，"
        "支持实时流数据处理和批量数据分析。"
    )

    # Para 9-11: 空段落
    doc.add_paragraph("")  # Para 9
    doc.add_paragraph("")  # Para 10
    doc.add_paragraph("")  # Para 11

    # Para 12: 预算表标题
    doc.add_paragraph("表1：预算分配")

    # 表格: 3行x3列 (不含表头，实际 4行x3列)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"

    # 表头
    headers = ["项目", "金额(万元)", "说明"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # 数据行
    data = [
        ["硬件设备", "120", "传感器、网关、服务器"],
        ["软件开发", "80", "平台开发、算法设计"],
        ["运维保障", "30", "人员培训、技术支持"],
    ]
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = val

    # Para 13-14: 空段落
    doc.add_paragraph("")  # Para 13
    doc.add_paragraph("")  # Para 14

    # Para 15: 时间计划
    doc.add_paragraph(
        "时间计划：项目周期为24个月，分为需求分析（第1-3月）、"
        "系统设计（第4-6月）、开发实施（第7-18月）、"
        "测试验收（第19-22月）和运维交接（第23-24月）五个阶段。"
    )

    # Para 16-19: 空段落
    doc.add_paragraph("")  # Para 16
    doc.add_paragraph("")  # Para 17
    doc.add_paragraph("")  # Para 18
    doc.add_paragraph("")  # Para 19

    # Para 20: 项目团队 (含 "Tiger公司" 和 "项目组")
    doc.add_paragraph(
        "项目团队：本项目由Tiger公司牵头，项目组成员包括高级工程师5名、"
        "数据分析师3名、测试工程师2名。项目组将定期召开技术评审会议，"
        "确保项目按时高质量交付。"
    )

    # Para 21-24: 空段落
    doc.add_paragraph("")  # Para 21
    doc.add_paragraph("")  # Para 22
    doc.add_paragraph("")  # Para 23
    doc.add_paragraph("")  # Para 24

    # Para 25: 参考文献
    doc.add_paragraph(
        "参考文献：\n"
        "[1] Wang et al., 2023, Smart Water Monitoring Systems\n"
        "[2] Liu et al., 2024, IoT-based Environmental Sensing\n"
        "[3] Chen et al., 2022, Big Data Analytics for Water Quality"
    )

    doc.save(output_path)
    print(f"✅ 已生成: {output_path}")

    # 验证段落索引
    verify_doc = Document(output_path)
    print(f"\n段落验证 (共{len(verify_doc.paragraphs)}段):")
    key_paras = [0, 3, 8, 12, 15, 20, 25]
    for i in key_paras:
        if i < len(verify_doc.paragraphs):
            text = verify_doc.paragraphs[i].text[:60]
            print(f"  Para {i}: {text}")
        else:
            print(f"  Para {i}: ⚠️ 超出范围")

    # 验证表格
    print(f"\n表格验证:")
    for t_idx, table in enumerate(verify_doc.tables):
        print(f"  表格 {t_idx}: {len(table.rows)}行 x {len(table.columns)}列")
        for r_idx, row in enumerate(table.rows):
            vals = [cell.text for cell in row.cells]
            print(f"    行{r_idx}: {vals}")

    return output_path


if __name__ == "__main__":
    output = sys.argv[1] if len(sys.argv) > 1 else None
    create_original_doc(output)

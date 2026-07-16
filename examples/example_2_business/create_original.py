#!/usr/bin/env python3
"""生成 Example 2: 商业报告的原始 docx 文档

段落结构 (表格不计入段落索引):
  Para 0:   标题 "Q3 2024 Sales Report"
  Para 1-2: 空段落
  Para 3:   Executive summary
  Para 4-7: 空段落
  Para 8:   表格标题 "Table 1: Regional Sales Performance"
  [表格: 6行x4列，含表头]
  Para 9-10: 空段落
  Para 11:  分析段落 (含 "2023年" 和 "$1.2M")
  Para 12-13: 空段落
  Para 14:  建议段落
"""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("需要安装 python-docx: pip install python-docx")
    sys.exit(1)


def create_original_doc(output_path: str = None):
    """生成原始商业报告 docx"""
    if output_path is None:
        output_path = str(Path(__file__).parent / "example_2_original.docx")

    doc = Document()

    # Para 0: 标题
    doc.add_heading("Q3 2024 Sales Report", level=0)

    # Para 1-2: 空段落
    doc.add_paragraph("")  # Para 1
    doc.add_paragraph("")  # Para 2

    # Para 3: Executive summary
    doc.add_paragraph(
        "This report provides a comprehensive overview of Q3 2024 sales performance "
        "across all regional divisions. Total revenue reached $4.8M, representing a "
        "15% increase over the previous quarter."
    )

    # Para 4-7: 空段落
    doc.add_paragraph("")  # Para 4
    doc.add_paragraph("")  # Para 5
    doc.add_paragraph("")  # Para 6
    doc.add_paragraph("")  # Para 7

    # Para 8: 表格标题
    doc.add_paragraph("Table 1: Regional Sales Performance")

    # Para 9: 表格 (6行含表头 = header + 5 data rows)
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"

    # 表头
    headers = ["Region", "Q3 Revenue", "Q2 Revenue", "Growth (%)"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # 数据行
    data = [
        ["North America", "$1.8M", "$1.5M", "20%"],
        ["Europe", "$1.2M", "$1.1M", "9%"],
        ["Asia Pacific", "$1.0M", "$0.8M", "25%"],
        ["Latin America", "$0.5M", "$0.4M", "25%"],
        ["Middle East", "$0.3M", "$0.3M", "0%"],
    ]
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = val

    # Para 10-11: 空段落
    doc.add_paragraph("")  # Para 10
    doc.add_paragraph("")  # Para 11

    # Para 12: 分析段落 (含 "2023年" 和 "$1.2M")
    doc.add_paragraph(
        "Compared to the same period in 2023年, overall performance has improved "
        "significantly. The Asia Pacific region contributed $1.2M in revenue, "
        "making it the fastest-growing market for the company."
    )

    # Para 13-14: 空段落
    doc.add_paragraph("")  # Para 13
    doc.add_paragraph("")  # Para 14

    # Para 15: 建议段落
    doc.add_paragraph(
        "Recommendations: Increase marketing spend in Europe to accelerate growth, "
        "and consider expanding the Asia Pacific sales team to sustain momentum."
    )

    doc.save(output_path)
    print(f"✅ 已生成: {output_path}")

    # 验证段落索引
    verify_doc = Document(output_path)
    print(f"\n段落验证 (共{len(verify_doc.paragraphs)}段):")
    for i in [0, 3, 8, 11, 14]:
        if i < len(verify_doc.paragraphs):
            text = verify_doc.paragraphs[i].text[:60]
            print(f"  Para {i}: {text}")
        else:
            print(f"  Para {i}: ⚠️ 超出范围")

    # 验证表格
    print(f"\n表格验证:")
    for t_idx, table in enumerate(verify_doc.tables):
        print(f"  表格 {t_idx}: {len(table.rows)}行 x {len(table.columns)}列")
        for r_idx, row in enumerate(table.rows):
            vals = [cell.text for cell in row.cells]
            print(f"    行{r_idx}: {vals}")

    return output_path


if __name__ == "__main__":
    output = sys.argv[1] if len(sys.argv) > 1 else None
    create_original_doc(output)

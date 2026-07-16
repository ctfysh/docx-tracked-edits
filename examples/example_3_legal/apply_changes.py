#!/usr/bin/env python3
"""将修订应用到 Example 3 的法律文档

工作流程:
  1. 用 md_to_json.py 将 changes.md 转换为 JSON
  2. 用 docx_revision 将 JSON 配置应用到原始 docx
"""
import json
import subprocess
import sys
from pathlib import Path

# 项目根目录
PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
SCRIPTS_DIR = PROJECT_ROOT / "scripts"
EXAMPLE_DIR = Path(__file__).parent


def main():
    md_file = EXAMPLE_DIR / "changes.md"
    json_file = EXAMPLE_DIR / "changes.json"

    # Step 1: 转换 MD → JSON
    print("📝 转换 changes.md → JSON...")
    result = subprocess.run(
        [sys.executable, str(SCRIPTS_DIR / "md_to_json.py"),
         str(md_file), str(json_file)],
        capture_output=True, text=True,
        cwd=str(EXAMPLE_DIR),
    )
    if result.returncode != 0:
        print(f"❌ MD 转换失败:\n{result.stderr}")
        sys.exit(1)
    print(result.stdout)

    # Step 2: 应用修订
    print("🔧 应用修订到 docx...")
    sys.path.insert(0, str(SCRIPTS_DIR))
    from docx_revision import ComprehensiveDocxReviewer

    with open(json_file, encoding="utf-8") as f:
        config = json.load(f)

    # 将全局 author 填充到 null 字段
    global_author = config.get("author", "Reviewer")
    for comment in config.get("comments", []):
        if comment.get("author") is None:
            comment["author"] = global_author
    for mod in config.get("text_modifications", []):
        if mod.get("author") is None:
            mod["author"] = global_author
    for mod in config.get("format_modifications", []):
        if mod.get("author") is None:
            mod["author"] = global_author
    for mod in config.get("table_modifications", []):
        if mod.get("author") is None:
            mod["author"] = global_author
    for mod in config.get("style_modifications", []):
        if mod.get("author") is None:
            mod["author"] = global_author

    # 解析相对路径为绝对路径
    source = Path(config["source"])
    if not source.is_absolute():
        source = EXAMPLE_DIR / source
    config["source"] = str(source)

    output = Path(config["output"])
    if not output.is_absolute():
        output = EXAMPLE_DIR / output
    config["output"] = str(output)

    reviewer = ComprehensiveDocxReviewer(config["source"])
    reviewer.apply_json_config(config)
    output_path = reviewer.save(config["output"])

    print(f"✅ 已生成修订文档: {output_path}")

    # 验证
    from docx import Document
    doc = Document(output_path)
    print(f"\n修订后段落预览 (共{len(doc.paragraphs)}段):")
    for i in [0, 3, 8, 12, 15, 18, 20]:
        if i < len(doc.paragraphs):
            text = doc.paragraphs[i].text[:80]
            print(f"  Para {i}: {text}")
        else:
            print(f"  Para {i}: ⚠️ 超出范围")


if __name__ == "__main__":
    main()

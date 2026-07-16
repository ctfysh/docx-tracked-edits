#!/usr/bin/env python3
"""生成 Example 3: 法律文档的原始 docx 文档

段落结构:
  Para 0:   标题 "Service Agreement"
  Para 1-2: 空段落
  Para 3:   合同双方 (甲方: 北京科技有限公司, 乙方: 张伟)
  Para 4-7: 空段落
  Para 8:   合同期限条款 (含日期 "2023年12月31日")
  Para 9-11: 空段落
  Para 12:  付款条款 (含金额 "伍万元整")
  Para 13-14: 空段落
  Para 15:  争议解决条款 ("如发生争议，双方应友好协商")
  Para 16-17: 空段落
  Para 18:  不可抗力条款
  Para 19:  空段落
  Para 20:  适用法律条款
"""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("需要安装 python-docx: pip install python-docx")
    sys.exit(1)


def create_original_doc(output_path: str = None):
    """生成原始法律文档 docx"""
    if output_path is None:
        output_path = str(Path(__file__).parent / "example_3_original.docx")

    doc = Document()

    # Para 0: 标题
    doc.add_heading("Service Agreement", level=0)

    # Para 1-2: 空段落
    doc.add_paragraph("")  # Para 1
    doc.add_paragraph("")  # Para 2

    # Para 3: 合同双方
    doc.add_paragraph(
        "甲方：北京科技有限公司（以下简称\u201c甲方\u201d）\n"
        "乙方：张伟（以下简称\u201c乙方\u201d）\n\n"
        "根据《中华人民共和国民法典》及相关法律法规的规定，"
        "甲乙双方在平等、自愿、公平、诚实信用的原则基础上，"
        "经协商一致，就甲方委托乙方提供技术服务事宜，订立本合同。"
    )

    # Para 4-7: 空段落
    doc.add_paragraph("")  # Para 4
    doc.add_paragraph("")  # Para 5
    doc.add_paragraph("")  # Para 6
    doc.add_paragraph("")  # Para 7

    # Para 8: 合同期限条款 (含旧日期)
    doc.add_paragraph(
        "第一条 合同期限\n"
        "本合同有效期自2023年1月1日起至2023年12月31日止。"
        "合同期满前30日内，任何一方未提出书面异议的，本合同自动续期一年。"
    )

    # Para 9-11: 空段落
    doc.add_paragraph("")  # Para 9
    doc.add_paragraph("")  # Para 10
    doc.add_paragraph("")  # Para 11

    # Para 12: 付款条款 (含旧金额)
    doc.add_paragraph(
        "第二条 付款条款\n"
        "甲方应在合同签订后十五个工作日内向乙方支付服务费人民币伍万元整（¥50,000.00）。"
        "逾期付款的，每逾期一日，甲方应按未付金额的万分之五向乙方支付违约金。"
    )

    # Para 13-14: 空段落
    doc.add_paragraph("")  # Para 13
    doc.add_paragraph("")  # Para 14

    # Para 15: 争议解决条款 (待标记为有风险)
    doc.add_paragraph(
        "第三条 争议解决\n"
        "如发生争议，双方应友好协商解决。协商不成的，任何一方均可向甲方所在地人民法院提起诉讼。"
    )

    # Para 16-17: 空段落
    doc.add_paragraph("")  # Para 16
    doc.add_paragraph("")  # Para 17

    # Para 18: 不可抗力条款 (待删除)
    doc.add_paragraph(
        "第四条 不可抗力\n"
        "因不可抗力导致本合同无法履行的，双方均不承担违约责任。"
        "不可抗力是指不能预见、不能避免且不能克服的客观情况，"
        "包括但不限于自然灾害、战争、政府行为等。"
    )

    # Para 19: 空段落
    doc.add_paragraph("")  # Para 19

    # Para 20: 适用法律条款
    doc.add_paragraph(
        "第五条 适用法律\n"
        "本合同的订立、效力、解释、履行及争议解决均适用中华人民共和国法律。"
    )

    doc.save(output_path)
    print(f"✅ 已生成: {output_path}")

    # 验证段落索引
    verify_doc = Document(output_path)
    print(f"\n段落验证 (共{len(verify_doc.paragraphs)}段):")
    for i in [0, 3, 8, 12, 15, 18, 20]:
        if i < len(verify_doc.paragraphs):
            text = verify_doc.paragraphs[i].text[:60]
            print(f"  Para {i}: {text}")
        else:
            print(f"  Para {i}: ⚠️ 超出范围")

    return output_path


if __name__ == "__main__":
    output = sys.argv[1] if len(sys.argv) > 1 else None
    create_original_doc(output)

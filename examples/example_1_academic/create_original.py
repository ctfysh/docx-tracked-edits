#!/usr/bin/env python3
"""生成 Example 1: 学术论文编辑的原始 docx 文档

段落结构:
  Para 0:  标题 "Research Paper on Flood Monitoring"
  Para 1-4: 空段落/填充段落
  Para 5:  摘要段落
  Para 6-7: 空段落
  Para 8:  需要修改的段落 (novel→improved, monitoring→detection, demonstrates→shows)
  Para 9-11: 空段落
  Para 12: 需要删除冗余的段落 (As previously reported...)
  Para 13-14: 空段落
  Para 15: 需要修复重复词的段落 (the the→the)
  Para 16-19: 空段落
  Para 20: 表格说明段落
"""
import sys
from pathlib import Path

# 确保可以导入 python-docx
try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("需要安装 python-docx: pip install python-docx")
    sys.exit(1)


def create_original_doc(output_path: str = None):
    """生成原始学术论文 docx"""
    if output_path is None:
        output_path = str(Path(__file__).parent / "example_1_original.docx")

    doc = Document()

    # Para 0: 标题
    doc.add_heading("Research Paper on Flood Monitoring", level=0)

    # Para 1-4: 填充段落
    doc.add_paragraph("Author: John Smith, Jane Doe")
    doc.add_paragraph("Department of Environmental Science, University of Example")
    doc.add_paragraph("Published: Journal of Hydrology, 2024")
    doc.add_paragraph("")  # Para 4: 空段落

    # Para 5: 摘要段落
    doc.add_paragraph(
        "Abstract: This paper presents a comprehensive analysis of flood monitoring "
        "techniques using remote sensing and machine learning approaches. We evaluate "
        "the performance of various models across different geographic regions and "
        "weather conditions."
    )

    # Para 6-7: 空段落
    doc.add_paragraph("")  # Para 6
    doc.add_paragraph("")  # Para 7

    # Para 8: 需要修改的段落
    doc.add_paragraph(
        "The novel approach for flood monitoring method demonstrates "
        "significant improvements in accuracy compared to traditional techniques."
    )

    # Para 9-11: 空段落
    doc.add_paragraph("")  # Para 9
    doc.add_paragraph("")  # Para 10
    doc.add_paragraph("")  # Para 11

    # Para 12: 需要删除冗余的段落
    doc.add_paragraph(
        "As previously reported in our earlier studies, the results show "
        "significant correlation between precipitation and flooding events."
    )

    # Para 13-14: 空段落
    doc.add_paragraph("")  # Para 13
    doc.add_paragraph("")  # Para 14

    # Para 15: 需要修复重复词的段落
    doc.add_paragraph(
        "The the methodology was validated using field data collected "
        "from multiple monitoring stations."
    )

    # Para 16-19: 空段落
    doc.add_paragraph("")  # Para 16
    doc.add_paragraph("")  # Para 17
    doc.add_paragraph("")  # Para 18
    doc.add_paragraph("")  # Para 19

    # Para 20: 表格说明段落
    doc.add_paragraph(
        "Table 1 summarizes the key findings from the analysis."
    )

    doc.save(output_path)
    print(f"✅ 已生成: {output_path}")

    # 验证段落索引
    verify_doc = Document(output_path)
    print(f"\n段落验证 (共{len(verify_doc.paragraphs)}段):")
    for i in [0, 5, 8, 12, 15, 20]:
        if i < len(verify_doc.paragraphs):
            text = verify_doc.paragraphs[i].text[:60]
            print(f"  Para {i}: {text}")
        else:
            print(f"  Para {i}: ⚠️ 超出范围")

    return output_path


if __name__ == "__main__":
    output = sys.argv[1] if len(sys.argv) > 1 else None
    create_original_doc(output)
